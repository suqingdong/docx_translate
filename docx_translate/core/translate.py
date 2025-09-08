import loguru
from docx import Document

from .translator import GoogleTranslator


class DocxTranslator:
    def __init__(self, source: str = 'auto', target: str = 'en'):
        """
        :param source: source language code, default is "auto"
        :param target: target language code, default is "en"
        """
        self.logger = loguru.logger
        self.translator = GoogleTranslator(source=source, target=target)

    def translate_docx(self, input_path: str, output_path: str):
        """
        :param input_path: input file path
        :param output_path: output file path
        """
        doc = Document(input_path)

        # translate header
        for section in doc.sections:
            if section.header.is_linked_to_previous:
                continue
            for paragraph in section.header.paragraphs:
                if paragraph.text.strip():
                    self.logger.debug(f'>>> translating header: {paragraph.text}')
                    translated_text = self.translator.translate(paragraph.text)
                    paragraph.text = translated_text

        # translate footer
        for section in doc.sections:
            if section.footer.is_linked_to_previous:
                continue
            for paragraph in section.footer.paragraphs:
                if paragraph.text.strip():
                    self.logger.debug(f'>>> translating footer: {paragraph.text}')
                    translated_text = self.translator.translate(paragraph.text)
                    paragraph.text = translated_text

        # translate paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                self.logger.debug(f'>>> translating paragraph: {paragraph.text}')
                translated_text = self.translator.translate(paragraph.text)
                paragraph.text = translated_text

        # translate tables
        for n, table in enumerate(doc.tables, 1):
            self.logger.debug(f'>>> translating table {n}')
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        translated_text = self.translator.translate(cell.text)
                        cell.text = translated_text

        # save translated docx
        doc.save(output_path)
        self.logger.info(f'>>> translated docx saved to {output_path}')